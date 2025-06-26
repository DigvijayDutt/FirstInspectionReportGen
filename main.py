# Import pandas for reading Excel data
import pandas as pd

# Import Document to create/edit Word files, and Inches for image sizing
from docx import Document
from docx.shared import Inches

# Read data from an Excel file named 'data.xls' into a DataFrame
DataFrame = pd.read_excel("./data.xls")

# Convert the DataFrame to a 2D NumPy array for easier row-by-row access
df = DataFrame.values

# Extract column headers from the Excel sheet (used as section titles in the document)
columns = DataFrame.columns

# Define a list of room names — each will have its own section and photos in the Word document
rooms = ["LIVING ROOM", "BEDROOM", "KITCHEN", "STORAGE"]

# Loop through each row (i.e., each claim record) in the Excel file
for j in range(len(df)):

    # Create a new Word document for the current claim
    doc = Document()

    # Add a company logo to the top of the document
    doc.add_picture("./logo.jpg", width=Inches(1.25))

    # Add a company website below the logo
    doc.add_paragraph("www.trinitycontents.com")

    # Add the main report title
    doc.add_heading("FIRST INSPECTION REPORT", level=1)

    # Insert a home image to make the document visually informative
    doc.add_picture("./images/home.jpg", width=Inches(2.5))

    # Loop through each column (field) from the Excel data
    for i in range(len(columns)):
        # Add the column title as a heading (e.g., "INSURER", "CLAIM #", etc.)
        doc.add_heading(f"{columns[i]}", level=2)

        # Add the corresponding cell data from the current row (claim)
        doc.add_paragraph(f"{df[j][i]}")

    # Start the photograph section
    doc.add_heading("PHOTOGRAPHS", level=2)

    # Loop through each predefined room
    for k in range(len(rooms)):
        # Add room name as a sub-heading
        doc.add_heading(rooms[k], level=3)

        # Add 4 photos for each room from the folder "./images/ROOM_NAME/"
        for l in range(4):
            try:
                # Insert each image into the document
                doc.add_picture(f"./images/{rooms[k]}/{l + 1}.jpg")
            except Exception as e:
                # In case image is missing or there's an error, skip it and optionally log
                print(f"Warning: Could not add image ./images/{rooms[k]}/{l + 1}.jpg - {e}")

    # Save the document with a unique name for each record
    doc.save(f"new{j}.docx")
